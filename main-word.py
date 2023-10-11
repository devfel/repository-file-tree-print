import os
import requests
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt

# Load variables from .env
load_dotenv()

OWNER = os.getenv("OWNER")
REPO = os.getenv("REPO")
BRANCH_SHA = os.getenv("BRANCH_SHA")
ACCESS_TOKEN = os.getenv("ACCESS_TOKEN")

# Fetch repository content
url = f"https://api.github.com/repos/{OWNER}/{REPO}/git/trees/{BRANCH_SHA}?recursive=1"
headers = {"Authorization": f"token {ACCESS_TOKEN}"}
response = requests.get(url, headers=headers)
data = response.json()


def get_icon_for_file(filename):
    extension = os.path.splitext(filename)[1]
    icon_mappings = {
        ".txt": "📜",
        ".md": "📜",
        ".js": "📜",
        ".ts": "📜",
        ".tsx": "📜",
        ".xml": "📜",
        ".json": "📜",
        ".jpg": "🖼️",
        ".jpeg": "🖼️",
        ".png": "🖼️",
        ".svg": "🖼️",
        ".ico": "🖼️",
        ".lock": "🔒",
    }
    return icon_mappings.get(extension, "📄")


# Create a new Word document
doc = Document()

# Process and construct Word content
for item in data["tree"]:
    depth = item["path"].count("/")
    indent = "  " * depth
    if item["type"] == "blob":
        icon = get_icon_for_file(item["path"])
        p = doc.add_paragraph(f"{indent}{icon} {item['path'].split('/')[-1]}")
    else:  # it's a directory
        icon = "📂"
        p = doc.add_paragraph(f"{indent}{icon} {item['path'].split('/')[-1]}")

    # Set the font to 'Segoe UI Emoji' for the paragraph
    for run in p.runs:
        run.font.name = "Segoe UI Emoji"
        run.font.size = Pt(12)

# Save to Word file
doc.save("repo_structure.docx")
